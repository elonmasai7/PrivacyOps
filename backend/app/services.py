import csv
import hashlib
import json
from datetime import datetime
from pathlib import Path

import boto3
import httpx
from botocore.exceptions import BotoCoreError, ClientError
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from sqlalchemy.orm import Session

from app.config import settings
from app.models import AuditLog, Integration, SecurityFinding


def write_audit_log(
    db: Session,
    *,
    organization_id: str | None,
    actor_user_id: str | None,
    action: str,
    target_type: str,
    target_id: str,
    metadata: dict | None = None,
) -> None:
    log = AuditLog(
        organization_id=organization_id,
        actor_user_id=actor_user_id,
        action=action,
        target_type=target_type,
        target_id=target_id,
        event_metadata=metadata or {},
    )
    db.add(log)
    db.commit()


def calculate_trust_readiness(answers: dict) -> dict:
    weights = {
        "handles_personal_data": 10,
        "handles_sensitive_data": 12,
        "serves_eu_users": 10,
        "has_privacy_policy": 8,
        "has_dpo": 8,
        "has_security_policies": 10,
        "has_incident_response": 10,
        "uses_cloud_services": 7,
        "uses_third_party_vendors": 7,
        "soc2_or_iso_required": 8,
        "processes_childrens_data": 10,
    }

    positive_score = 0
    risk_areas: list[str] = []
    for key, weight in weights.items():
        value = bool(answers.get(key, False))
        if key in {
            "has_privacy_policy",
            "has_dpo",
            "has_security_policies",
            "has_incident_response",
        }:
            if value:
                positive_score += weight
            else:
                risk_areas.append(key)
        else:
            if value:
                positive_score += int(weight * 0.5)
                if key in {"handles_sensitive_data", "processes_childrens_data"}:
                    risk_areas.append(key)
            else:
                positive_score += weight

    suggested_frameworks = ["Kenya Data Protection Act", "GDPR"]
    if answers.get("soc2_or_iso_required"):
        suggested_frameworks.extend(["SOC 2 Readiness", "ISO 27001 Readiness"])

    required_workflows = [
        "Data inventory",
        "RoPA",
        "Evidence vault",
        "Incident and breach workflow",
        "Vendor risk review",
    ]
    missing_evidence = [
        "Privacy policy" if not answers.get("has_privacy_policy") else "",
        "Security policy" if not answers.get("has_security_policies") else "",
        "Incident response procedure" if not answers.get("has_incident_response") else "",
    ]

    score = max(0, min(100, positive_score))
    return {
        "trust_readiness_score": score,
        "suggested_frameworks": suggested_frameworks,
        "required_workflows": required_workflows,
        "risk_areas": risk_areas,
        "recommended_next_actions": [
            "Complete onboarding evidence checklist",
            "Assign control owners",
            "Connect at least one integration for automation",
        ],
        "missing_evidence": [item for item in missing_evidence if item],
    }


def hash_file(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def ensure_dir(path: str) -> Path:
    p = Path(path)
    p.mkdir(parents=True, exist_ok=True)
    return p


def generate_report_files(base_dir: str, report_id: str, payload: dict) -> dict[str, str]:
    export_dir = ensure_dir(base_dir)

    json_path = export_dir / f"{report_id}.json"
    csv_path = export_dir / f"{report_id}.csv"
    docx_path = export_dir / f"{report_id}.docx"
    pdf_path = export_dir / f"{report_id}.pdf"

    json_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")

    with csv_path.open("w", encoding="utf-8", newline="") as fp:
        writer = csv.writer(fp)
        writer.writerow(["key", "value"])
        for key, value in payload.items():
            writer.writerow([key, json.dumps(value)])

    doc = Document()
    doc.add_heading("PrivacyOps Africa Report", level=1)
    for key, value in payload.items():
        doc.add_paragraph(f"{key}: {json.dumps(value)}")
    doc.save(docx_path)

    c = canvas.Canvas(str(pdf_path), pagesize=A4)
    y = 800
    c.setFont("Helvetica", 12)
    c.drawString(40, y, "PrivacyOps Africa Report")
    y -= 30
    for key, value in payload.items():
        text = f"{key}: {json.dumps(value)}"
        c.drawString(40, y, text[:110])
        y -= 18
        if y < 80:
            c.showPage()
            y = 800
    c.save()

    return {
        "json": str(json_path),
        "csv": str(csv_path),
        "docx": str(docx_path),
        "pdf": str(pdf_path),
    }


def ai_guardrailed_response(prompt: str, context: dict) -> dict:
    legal_keywords = ["lawful basis", "contract", "regulator", "legal", "breach"]
    requires_legal_review = any(keyword in prompt.lower() for keyword in legal_keywords)

    guidance = (
        "This guidance is generated for compliance operations and does not replace legal advice. "
        "Use the listed controls, collected evidence, and your internal counsel for final decisions."
    )
    if context.get("gap_summary"):
        guidance += f" Gap focus: {context['gap_summary']}."

    return {
        "guidance": guidance,
        "confidence_level": "medium",
        "requires_legal_review": requires_legal_review,
        "source_references": context.get("source_references", []),
        "next_action": "Assign a compliance owner, attach evidence, and route legal-sensitive items for counsel review.",
    }


async def scan_github(integration: Integration, db: Session) -> list[SecurityFinding]:
    token = integration.config.get("personal_access_token")
    headers = {
        "Accept": "application/vnd.github+json",
        "Authorization": f"Bearer {token}",
        "X-GitHub-Api-Version": "2022-11-28",
    }

    findings: list[SecurityFinding] = []
    async with httpx.AsyncClient(timeout=30) as client:
        repos_response = await client.get(f"{settings.github_api_base}/user/repos", headers=headers)
        repos_response.raise_for_status()
        repos = repos_response.json()

        for repo in repos[:50]:
            if repo.get("private") is False:
                findings.append(
                    SecurityFinding(
                        organization_id=integration.organization_id,
                        integration_id=integration.id,
                        title=f"Public repository detected: {repo['full_name']}",
                        severity="medium",
                        category="repository_visibility",
                        details={"repo": repo["full_name"], "visibility": "public"},
                    )
                )

            repo_name = repo["name"]
            owner = repo["owner"]["login"]
            branch_resp = await client.get(
                f"{settings.github_api_base}/repos/{owner}/{repo_name}/branches/{repo['default_branch']}",
                headers=headers,
            )
            if branch_resp.status_code == 200:
                protection = branch_resp.json().get("protected", False)
                if not protection:
                    findings.append(
                        SecurityFinding(
                            organization_id=integration.organization_id,
                            integration_id=integration.id,
                            title=f"Branch protection missing: {repo['full_name']}",
                            severity="high",
                            category="branch_protection",
                            details={"repo": repo["full_name"], "default_branch": repo["default_branch"]},
                        )
                    )

    for finding in findings:
        db.add(finding)
    integration.last_synced_at = datetime.utcnow()
    integration.status = "connected"
    integration.last_error = None
    db.add(integration)
    db.commit()
    return findings


async def scan_gitlab(integration: Integration, db: Session) -> list[SecurityFinding]:
    token = integration.config.get("personal_access_token")
    headers = {"PRIVATE-TOKEN": token}

    findings: list[SecurityFinding] = []
    async with httpx.AsyncClient(timeout=30) as client:
        projects_response = await client.get(
            f"{settings.gitlab_api_base}/projects",
            headers=headers,
            params={"membership": "true", "simple": "true", "per_page": 100, "archived": "false"},
        )
        projects_response.raise_for_status()
        projects = projects_response.json()

        for project in projects[:100]:
            project_id = project.get("id")
            project_name = project.get("path_with_namespace", str(project_id))
            default_branch = project.get("default_branch")

            if project.get("visibility") == "public":
                findings.append(
                    SecurityFinding(
                        organization_id=integration.organization_id,
                        integration_id=integration.id,
                        title=f"Public project detected: {project_name}",
                        severity="medium",
                        category="repository_visibility",
                        details={"project": project_name, "visibility": "public"},
                    )
                )

            if default_branch:
                protected_resp = await client.get(
                    f"{settings.gitlab_api_base}/projects/{project_id}/protected_branches/{default_branch}",
                    headers=headers,
                )
                if protected_resp.status_code == 404:
                    findings.append(
                        SecurityFinding(
                            organization_id=integration.organization_id,
                            integration_id=integration.id,
                            title=f"Protected branch missing: {project_name}",
                            severity="high",
                            category="branch_protection",
                            details={"project": project_name, "default_branch": default_branch},
                        )
                    )

            approvals_resp = await client.get(
                f"{settings.gitlab_api_base}/projects/{project_id}/approvals",
                headers=headers,
            )
            if approvals_resp.status_code == 200:
                approvals_before_merge = approvals_resp.json().get("approvals_before_merge", 0)
                if approvals_before_merge == 0:
                    findings.append(
                        SecurityFinding(
                            organization_id=integration.organization_id,
                            integration_id=integration.id,
                            title=f"Required merge approvals missing: {project_name}",
                            severity="medium",
                            category="merge_approval_policy",
                            details={"project": project_name, "approvals_before_merge": approvals_before_merge},
                        )
                    )

    for finding in findings:
        db.add(finding)
    integration.last_synced_at = datetime.utcnow()
    integration.status = "connected"
    integration.last_error = None
    db.add(integration)
    db.commit()
    return findings


def validate_aws_credentials(config: dict) -> dict:
    session = boto3.session.Session(
        aws_access_key_id=config.get("aws_access_key_id"),
        aws_secret_access_key=config.get("aws_secret_access_key"),
        aws_session_token=config.get("aws_session_token"),
        region_name=config.get("aws_region") or settings.aws_region,
    )
    sts = session.client("sts")
    identity = sts.get_caller_identity()
    return {
        "account": identity.get("Account"),
        "arn": identity.get("Arn"),
    }


def scan_aws(integration: Integration, db: Session) -> list[SecurityFinding]:
    config = integration.config
    session = boto3.session.Session(
        aws_access_key_id=config.get("aws_access_key_id"),
        aws_secret_access_key=config.get("aws_secret_access_key"),
        aws_session_token=config.get("aws_session_token"),
        region_name=config.get("aws_region") or settings.aws_region,
    )

    findings: list[SecurityFinding] = []

    try:
        iam = session.client("iam")
        users = iam.list_users().get("Users", [])
        for user in users[:100]:
            mfa_devices = iam.list_mfa_devices(UserName=user["UserName"]).get("MFADevices", [])
            if not mfa_devices:
                findings.append(
                    SecurityFinding(
                        organization_id=integration.organization_id,
                        integration_id=integration.id,
                        title=f"IAM user without MFA: {user['UserName']}",
                        severity="high",
                        category="iam_mfa",
                        details={"user_name": user["UserName"]},
                    )
                )
    except (BotoCoreError, ClientError) as exc:
        findings.append(
            SecurityFinding(
                organization_id=integration.organization_id,
                integration_id=integration.id,
                title="Unable to enumerate IAM users",
                severity="medium",
                category="iam_access",
                details={"error": str(exc)},
            )
        )

    try:
        s3 = session.client("s3")
        buckets = s3.list_buckets().get("Buckets", [])
        for bucket in buckets[:100]:
            bucket_name = bucket["Name"]
            try:
                block = s3.get_public_access_block(Bucket=bucket_name)
                conf = block.get("PublicAccessBlockConfiguration", {})
                all_blocked = all(
                    [
                        conf.get("BlockPublicAcls", False),
                        conf.get("IgnorePublicAcls", False),
                        conf.get("BlockPublicPolicy", False),
                        conf.get("RestrictPublicBuckets", False),
                    ]
                )
                if not all_blocked:
                    findings.append(
                        SecurityFinding(
                            organization_id=integration.organization_id,
                            integration_id=integration.id,
                            title=f"S3 bucket may be publicly exposed: {bucket_name}",
                            severity="high",
                            category="s3_public_access",
                            details={"bucket": bucket_name, "public_access_block": conf},
                        )
                    )
            except (BotoCoreError, ClientError):
                findings.append(
                    SecurityFinding(
                        organization_id=integration.organization_id,
                        integration_id=integration.id,
                        title=f"S3 public access block missing or unreadable: {bucket_name}",
                        severity="medium",
                        category="s3_public_access",
                        details={"bucket": bucket_name},
                    )
                )

            try:
                s3.get_bucket_encryption(Bucket=bucket_name)
            except (BotoCoreError, ClientError):
                findings.append(
                    SecurityFinding(
                        organization_id=integration.organization_id,
                        integration_id=integration.id,
                        title=f"S3 bucket encryption not configured: {bucket_name}",
                        severity="medium",
                        category="s3_encryption",
                        details={"bucket": bucket_name},
                    )
                )
    except (BotoCoreError, ClientError) as exc:
        findings.append(
            SecurityFinding(
                organization_id=integration.organization_id,
                integration_id=integration.id,
                title="Unable to enumerate S3 buckets",
                severity="medium",
                category="s3_access",
                details={"error": str(exc)},
            )
        )

    try:
        ec2 = session.client("ec2")
        security_groups = ec2.describe_security_groups().get("SecurityGroups", [])
        for group in security_groups[:200]:
            for permission in group.get("IpPermissions", []):
                for ip_range in permission.get("IpRanges", []):
                    if ip_range.get("CidrIp") == "0.0.0.0/0":
                        from_port = permission.get("FromPort")
                        protocol = permission.get("IpProtocol")
                        high_risk = from_port in {22, 3389} or protocol == "-1"
                        if high_risk:
                            findings.append(
                                SecurityFinding(
                                    organization_id=integration.organization_id,
                                    integration_id=integration.id,
                                    title=f"Open security group rule: {group.get('GroupName')}",
                                    severity="high",
                                    category="security_group_exposure",
                                    details={
                                        "group_id": group.get("GroupId"),
                                        "group_name": group.get("GroupName"),
                                        "from_port": from_port,
                                        "protocol": protocol,
                                    },
                                )
                            )
    except (BotoCoreError, ClientError) as exc:
        findings.append(
            SecurityFinding(
                organization_id=integration.organization_id,
                integration_id=integration.id,
                title="Unable to enumerate security groups",
                severity="medium",
                category="ec2_security_groups",
                details={"error": str(exc)},
            )
        )

    try:
        cloudtrail = session.client("cloudtrail")
        trails = cloudtrail.describe_trails().get("trailList", [])
        if not trails:
            findings.append(
                SecurityFinding(
                    organization_id=integration.organization_id,
                    integration_id=integration.id,
                    title="CloudTrail appears not configured",
                    severity="high",
                    category="cloudtrail",
                    details={},
                )
            )
    except (BotoCoreError, ClientError) as exc:
        findings.append(
            SecurityFinding(
                organization_id=integration.organization_id,
                integration_id=integration.id,
                title="Unable to check CloudTrail status",
                severity="medium",
                category="cloudtrail",
                details={"error": str(exc)},
            )
        )

    for finding in findings:
        db.add(finding)
    integration.last_synced_at = datetime.utcnow()
    integration.status = "connected"
    integration.last_error = None
    db.add(integration)
    db.commit()
    return findings
